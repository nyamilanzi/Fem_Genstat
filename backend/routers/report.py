"""
Report generation router
"""

import os
import tempfile
import pandas as pd
from typing import Dict, Any
from fastapi import APIRouter, HTTPException, Depends, Request
from fastapi.responses import FileResponse, HTMLResponse
from jinja2 import Environment, FileSystemLoader
import weasyprint
from docx import Document
from docx.shared import Inches

from models.schemas import ReportRequest, ReportResponse, ErrorResponse
from services.cache import DataCache

router = APIRouter()

def get_data_cache(request: Request) -> DataCache:
    """Dependency to get data cache instance"""
    return request.app.state.data_cache

@router.post("/report", response_model=ReportResponse)
async def generate_report(
    request: ReportRequest,
    cache: DataCache = Depends(get_data_cache)
):
    """Generate HTML, PDF, and DOCX reports"""
    
    # Get session and analysis data
    session = cache.get_session(request.session_id)
    if not session:
        raise HTTPException(
            status_code=404,
            detail="Session not found or expired"
        )
    
    analysis_results = session.get("analysis_results")
    if not analysis_results:
        raise HTTPException(
            status_code=400,
            detail="No analysis results found. Please run analysis first."
        )
    
    # Generate HTML report
    html_url = await _generate_html_report(request, analysis_results, session)
    
    # Generate PDF report (optional - may fail)
    pdf_url = None
    try:
        pdf_url = await _generate_pdf_report(request, analysis_results, session)
    except Exception as e:
        print(f"PDF generation failed: {str(e)}")
        # Continue without PDF
    
    # Generate DOCX report (optional - may fail)
    docx_url = None
    try:
        docx_url = await _generate_docx_report(request, analysis_results, session)
    except Exception as e:
        print(f"DOCX generation failed: {str(e)}")
        # Continue without DOCX
    
    return ReportResponse(
        html_url=html_url,
        pdf_url=pdf_url or "",
        docx_url=docx_url
    )

async def _generate_html_report(
    request: ReportRequest, 
    analysis_results: Dict[str, Any], 
    session: Dict[str, Any]
) -> str:
    """Generate HTML report"""
    
    # Set up Jinja2 environment
    template_dir = os.path.join(os.path.dirname(__file__), "..", "templates")
    env = Environment(loader=FileSystemLoader(template_dir))
    template = env.get_template("report.html.j2")
    
    # Prepare template data
    template_data = {
        "title": request.title,
        "organization": request.organization or "Not specified",
        "authors": request.authors or ["Analysis Tool"],
        "notes": request.notes or "",
        "analysis_results": analysis_results,
        "session_metadata": session.get("metadata", {}),
        "generation_date": pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"),
        "suppress_threshold": analysis_results["settings"].get("suppress_threshold", 5)
    }
    
    # Render HTML
    html_content = template.render(**template_data)
    
    # Save HTML file
    static_dir = "static/reports"
    os.makedirs(static_dir, exist_ok=True)
    
    html_filename = f"{request.session_id}_report.html"
    html_path = os.path.join(static_dir, html_filename)
    
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    return f"/static/reports/{html_filename}"

async def _generate_pdf_report(
    request: ReportRequest,
    analysis_results: Dict[str, Any],
    session: Dict[str, Any]
) -> str:
    """Generate PDF report using WeasyPrint"""
    
    # First generate HTML
    html_url = await _generate_html_report(request, analysis_results, session)
    html_path = os.path.join("static", html_url.lstrip("/static/reports/"))
    
    # Ensure directory exists
    pdf_dir = "static/reports"
    os.makedirs(pdf_dir, exist_ok=True)
    
    # Convert to PDF
    pdf_filename = f"{request.session_id}_report.pdf"
    pdf_path = os.path.join(pdf_dir, pdf_filename)
    
    try:
        # Read HTML content and convert to PDF
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Use base_url for relative paths
        base_url = os.path.dirname(os.path.abspath(html_path))
        weasyprint.HTML(string=html_content, base_url=base_url).write_pdf(pdf_path)
        return f"/static/reports/{pdf_filename}"
    except Exception as e:
        print(f"PDF generation error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error generating PDF: {str(e)}"
        )

async def _generate_docx_report(
    request: ReportRequest,
    analysis_results: Dict[str, Any],
    session: Dict[str, Any]
) -> str:
    """Generate DOCX report"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading(request.title, 0)
    
    # FEMSTAT branding
    p = doc.add_paragraph()
    p.add_run('Developed by ').bold = True
    p.add_run('FEMSTAT from Femanalytica').bold = True
    p.add_run(' - Professional Gender Analysis Platform')
    
    # Metadata
    doc.add_heading('Report Information', level=1)
    p = doc.add_paragraph()
    p.add_run('Organization: ').bold = True
    p.add_run(request.organization or 'Not specified')
    
    p = doc.add_paragraph()
    p.add_run('Authors: ').bold = True
    p.add_run(', '.join(request.authors) if request.authors else 'Analysis Tool')
    
    p = doc.add_paragraph()
    p.add_run('Generated: ').bold = True
    p.add_run(pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S"))
    
    # Dataset information
    doc.add_heading('Dataset Information', level=1)
    metadata = session.get("metadata", {})
    p = doc.add_paragraph()
    p.add_run('Filename: ').bold = True
    p.add_run(metadata.get("filename", "Unknown"))
    
    p = doc.add_paragraph()
    p.add_run('Total rows: ').bold = True
    p.add_run(str(metadata.get("total_rows", "Unknown")))
    
    # Gender summary
    doc.add_heading('Gender Distribution', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Gender'
    hdr_cells[1].text = 'N'
    hdr_cells[2].text = 'Percent'
    hdr_cells[3].text = 'Missing %'
    
    for gender_summary in analysis_results.get("by_gender", []):
        row_cells = table.add_row().cells
        row_cells[0].text = str(gender_summary["gender"])
        row_cells[1].text = str(gender_summary["n"])
        row_cells[2].text = str(gender_summary["pct"])
        row_cells[3].text = str(gender_summary["missing_pct"])
    
    # Continuous variables
    if analysis_results.get("continuous"):
        doc.add_heading('Continuous Variables', level=1)
        
        for var_result in analysis_results["continuous"]:
            doc.add_heading(f'Variable: {var_result["var"]}', level=2)
            
            # Summary table
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            headers = ['Gender', 'N', 'Mean', 'SD', 'Median', 'IQR', 'Min', 'Max']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            for stat in var_result["table"]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(stat["gender"])
                row_cells[1].text = str(stat["n"])
                row_cells[2].text = str(stat["mean"])
                row_cells[3].text = str(stat["sd"])
                row_cells[4].text = str(stat["median"])
                row_cells[5].text = str(stat["iqr"])
                row_cells[6].text = str(stat["min"])
                row_cells[7].text = str(stat["max"])
            
            # Test results
            test = var_result.get("test", {})
            if test:
                p = doc.add_paragraph()
                p.add_run('Statistical Test: ').bold = True
                p.add_run(test.get("name", "Unknown"))
                
                p = doc.add_paragraph()
                p.add_run('P-value: ').bold = True
                p.add_run(str(test.get("p", "N/A")))
                
                if test.get("note"):
                    p = doc.add_paragraph()
                    p.add_run('Note: ').italic = True
                    p.add_run(test["note"])
    
    # Categorical variables
    if analysis_results.get("categorical"):
        doc.add_heading('Categorical Variables', level=1)
        
        for var_result in analysis_results["categorical"]:
            doc.add_heading(f'Variable: {var_result["var"]}', level=2)
            
            # Summary table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Level'
            hdr_cells[1].text = 'Gender'
            hdr_cells[2].text = 'Count (%)'
            
            for level in var_result["table"]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(level["level"])
                row_cells[1].text = str(level["gender"])
                # Format categorical data: integer for counts, one decimal for percentages
                n_val = level['n']
                if isinstance(n_val, (int, float)) and not isinstance(n_val, str):
                    n_val = int(n_val)
                pct_val = level['pct']
                if isinstance(pct_val, (int, float)) and not isinstance(pct_val, str):
                    pct_val = f"{pct_val:.1f}"
                row_cells[2].text = f"{n_val} ({pct_val}%)"
            
            # Test results
            test = var_result.get("test", {})
            if test:
                p = doc.add_paragraph()
                p.add_run('Statistical Test: ').bold = True
                p.add_run(test.get("name", "Unknown"))
                
                p = doc.add_paragraph()
                p.add_run('P-value: ').bold = True
                p.add_run(str(test.get("p", "N/A")))
    
    # Gender Bias Assessment
    if analysis_results.get("gender_bias"):
        doc.add_heading('Gender Bias Assessment', level=1)
        
        gender_bias = analysis_results["gender_bias"]
        
        # Overall Summary
        p = doc.add_paragraph()
        p.add_run('Overall Summary: ').bold = True
        p.add_run(gender_bias.get("overall_summary", ""))
        
        # Statistical Disparities
        if gender_bias.get("statistical_disparities"):
            doc.add_heading('Statistical Disparities', level=2)
            for disparity in gender_bias["statistical_disparities"]:
                p = doc.add_paragraph()
                p.add_run(f"{disparity['variable']} ({disparity['type']}): ").bold = True
                p.add_run(disparity.get("interpretation", ""))
        
        # Representation Gaps
        if gender_bias.get("representation_gaps"):
            doc.add_heading('Representation Gaps', level=2)
            for gap in gender_bias["representation_gaps"]:
                doc.add_paragraph(gap.get("interpretation", ""), style='List Bullet')
        
        # Missing Data Bias
        if gender_bias.get("missing_data_bias"):
            doc.add_heading('Missing Data Bias', level=2)
            for bias in gender_bias["missing_data_bias"]:
                doc.add_paragraph(bias.get("interpretation", ""), style='List Bullet')
        
        # Gender Transformative Insights
        if gender_bias.get("gender_transformative_insights"):
            doc.add_heading('Gender Transformative Insights', level=2)
            for insight in gender_bias["gender_transformative_insights"]:
                doc.add_paragraph(insight, style='List Bullet')
        
        # Recommendations
        if gender_bias.get("recommendations"):
            doc.add_heading('Recommendations', level=2)
            for recommendation in gender_bias["recommendations"]:
                doc.add_paragraph(recommendation, style='List Bullet')
    
    # Key Findings
    doc.add_heading('Key Findings and Interpretations', level=1)
    
    # Continuous variables interpretations
    if analysis_results.get("continuous"):
        doc.add_heading('Continuous Variables - Summary', level=2)
        for var_result in analysis_results["continuous"]:
            var_name = var_result["var"]
            test = var_result.get("test", {})
            p_val = test.get("p")
            
            p = doc.add_paragraph()
            p.add_run(f"{var_name}: ").bold = True
            if p_val and isinstance(p_val, (int, float)) and not isinstance(p_val, str):
                if p_val < 0.05:
                    p.add_run(f"Significant difference found (p={p_val:.4f}). There is a statistically significant difference in {var_name} between gender groups.")
                else:
                    p.add_run(f"No significant difference (p={p_val:.4f}). There is no statistically significant difference in {var_name} between gender groups.")
            else:
                p.add_run("Statistical test results are not available for this variable.")
    
    # Categorical variables interpretations
    if analysis_results.get("categorical"):
        doc.add_heading('Categorical Variables - Summary', level=2)
        for var_result in analysis_results["categorical"]:
            var_name = var_result["var"]
            test = var_result.get("test", {})
            p_val = test.get("p")
            
            p = doc.add_paragraph()
            p.add_run(f"{var_name}: ").bold = True
            if p_val and isinstance(p_val, (int, float)) and not isinstance(p_val, str):
                if p_val < 0.05:
                    p.add_run(f"Significant association found (p={p_val:.4f}). There is a statistically significant association between {var_name} and gender, indicating gender-based differences in distribution.")
                else:
                    p.add_run(f"No significant association (p={p_val:.4f}). There is no statistically significant association between {var_name} and gender.")
            else:
                p.add_run("Statistical test results are not available for this variable.")
    
    # Notes
    if request.notes:
        doc.add_heading('Notes', level=1)
        doc.add_paragraph(request.notes)
    
    # Footer
    p = doc.add_paragraph()
    p.add_run(f"Report generated by FEMSTAT (from Femanalytica) on {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    
    # Save DOCX file
    static_dir = "static/reports"
    os.makedirs(static_dir, exist_ok=True)
    
    docx_filename = f"{request.session_id}_report.docx"
    docx_path = os.path.join(static_dir, docx_filename)
    
    doc.save(docx_path)
    
    return f"/static/reports/{docx_filename}"

@router.get("/report/{session_id}")
async def view_report(session_id: str, cache: DataCache = Depends(get_data_cache)):
    """View HTML report"""
    
    session = cache.get_session(session_id)
    if not session:
        raise HTTPException(
            status_code=404,
            detail="Session not found or expired"
        )
    
    analysis_results = session.get("analysis_results")
    if not analysis_results:
        raise HTTPException(
            status_code=400,
            detail="No analysis results found"
        )
    
    # Generate HTML report
    request = ReportRequest(
        session_id=session_id,
        title="Gender Analysis Report"
    )
    
    html_url = await _generate_html_report(request, analysis_results, session)
    html_path = f"static{html_url}"
    
    return FileResponse(html_path, media_type="text/html")
