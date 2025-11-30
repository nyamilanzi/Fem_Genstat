import streamlit as st
import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
from scipy import stats
from scipy.stats import mannwhitneyu, chi2_contingency
import io
from datetime import datetime
import tempfile
import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors as rl_colors
from docx import Document
from docx.shared import Inches

# --- BRAND COLORS & LOGO ---
BRAND_COLORS = ['#0b71a1', '#8245aa', '#67b7d1', '#4f64ac']
ORG_NAME = "FemAnalytica"
LOGO_PATH = "femanalytica_logo.png"  # Save your provided logo as this file in the project root

st.set_page_config(
    page_title=f"{ORG_NAME} Gender Analysis Platform",
    page_icon="🌍",
    layout="wide"
)

# --- HEADER & LOGO ---
# Always use the provided FemAnalytica logo
st.image(LOGO_PATH, width=350)
st.markdown(f"<h1 style='color:{BRAND_COLORS[0]};'>{ORG_NAME} Gender Disaggregated Analysis Platform</h1>", unsafe_allow_html=True)
st.markdown(f"<h4 style='color:{BRAND_COLORS[2]};'>Upload your data and explore gender differences with professional, branded visualizations.</h4>", unsafe_allow_html=True)

# --- FILE UPLOAD ---
uploaded_file = st.sidebar.file_uploader("Upload your CSV dataset", type=["csv"])
data_dict_file = st.sidebar.file_uploader("(Optional) Upload a data dictionary CSV", type=["csv"], key="dict")

# No logo upload option; always use FemAnalytica logo
logo_file = LOGO_PATH

data_dict = None
if data_dict_file is not None:
    data_dict = pd.read_csv(data_dict_file)

if uploaded_file is not None:
    df = pd.read_csv(uploaded_file)
    st.success("File uploaded successfully!")
    st.write("Preview of your data:")
    st.dataframe(df.head())

    # --- VARIABLE SELECTION ---
    st.sidebar.header("Variable Selection")
    columns = df.columns.tolist()
    gender_col = st.sidebar.selectbox("Select the gender column", columns)
    outcome_col = st.sidebar.selectbox("Select the outcome variable", [col for col in columns if col != gender_col])
    group_col = st.sidebar.selectbox("(Optional) Group by another variable", ["None"] + [col for col in columns if col not in [gender_col, outcome_col]])
    plot_types = st.sidebar.multiselect(
        "Select plot types (choose one or more)",
        ["Violin Plot", "Box Plot", "Bar Plot", "Scatter Plot", "Pie Chart", "Correlation Heatmap"],
        default=["Violin Plot"]
    )
    report_format = st.sidebar.radio("Choose report format", ["PDF", "Word"])

    # --- GENDER MAPPING ---
    gender_values = df[gender_col].dropna().unique()
    gender_map = {}
    for val in gender_values:
        gender_map[val] = st.sidebar.selectbox(f"Map value '{val}' to:", ['Female', 'Male', 'Other'], key=f"gender_{val}")
    df['gender_analysis'] = df[gender_col].map(gender_map)
    gendered_df = df[df['gender_analysis'].isin(['Male', 'Female'])]

    # --- OUTCOME TYPE ---
    outcome_type = 'numeric' if np.issubdtype(df[outcome_col].dropna().dtype, np.number) else 'categorical'
    st.write(f"Outcome variable '{outcome_col}' detected as: {outcome_type}")
    if data_dict is not None and outcome_col in data_dict.columns:
        st.info(f"Description: {data_dict[outcome_col].values[0]}")

    # --- VISUALIZATION & NARRATIVE ---
    st.header(f"Gender Disaggregated Analysis for '{outcome_col}'")
    plot_narratives = []
    plot_images = []
    def plot_for_group(sub_df, group_val=None):
        for plot_type in plot_types:
            buf = io.BytesIO()
            group_label = f" ({group_col}: {group_val})" if group_val is not None else ""
            if plot_type == "Violin Plot" and outcome_type == 'numeric':
                fig, ax = plt.subplots()
                sns.violinplot(data=sub_df, x='gender_analysis', y=outcome_col, hue='gender_analysis', palette=BRAND_COLORS[:2], legend=False)
                plt.title(f"Violin Plot of {outcome_col} by Gender{group_label}")
                plt.xlabel('Gender')
                plt.ylabel(outcome_col)
                st.pyplot(fig)
                fig.savefig(buf, format='png')
                plot_images.append((f"Violin Plot of {outcome_col} by Gender{group_label}", buf.getvalue()))
                st.download_button(f"Download Violin Plot{group_label}", buf.getvalue(), file_name=f"violin_plot{group_label}.png", mime="image/png")
                plot_narratives.append(f"The violin plot shows the distribution and spread of {outcome_col} for males and females{group_label}. Look for differences in the width and center of the distributions to spot gender differences.")
            if plot_type == "Box Plot" and outcome_type == 'numeric':
                fig, ax = plt.subplots()
                sns.boxplot(data=sub_df, x='gender_analysis', y=outcome_col, hue='gender_analysis', palette=BRAND_COLORS[:2], legend=False)
                plt.title(f"Box Plot of {outcome_col} by Gender{group_label}")
                plt.xlabel('Gender')
                plt.ylabel(outcome_col)
                st.pyplot(fig)
                fig.savefig(buf, format='png')
                plot_images.append((f"Box Plot of {outcome_col} by Gender{group_label}", buf.getvalue()))
                plot_narratives.append(f"The box plot compares the medians and interquartile ranges of {outcome_col} by gender{group_label}. Differences in box heights or medians highlight gender disparities.")
            if plot_type == "Bar Plot" and outcome_type == 'categorical':
                fig, ax = plt.subplots()
                percent_tab = pd.crosstab(sub_df['gender_analysis'], sub_df[outcome_col], normalize='index') * 100
                n_bars = percent_tab.shape[1]
                percent_tab.T.plot(kind='bar', ax=ax, color=BRAND_COLORS[:n_bars])
                plt.title(f"{outcome_col} by Gender (Percentage){group_label}")
                plt.xlabel(outcome_col)
                plt.ylabel('Percentage')
                plt.legend(title='Gender')
                st.pyplot(fig)
                fig.savefig(buf, format='png')
                plot_images.append((f"Bar Plot of {outcome_col} by Gender{group_label}", buf.getvalue()))
                st.download_button(f"Download Bar Plot{group_label}", buf.getvalue(), file_name=f"bar_plot{group_label}.png", mime="image/png")
                plot_narratives.append(f"The bar plot shows the percentage of each {outcome_col} category for males and females{group_label}. Differences in bar heights indicate gender differences in categorical outcomes.")
            if plot_type == "Scatter Plot":
                numeric_cols = [col for col in columns if np.issubdtype(df[col].dropna().dtype, np.number)]
                if len(numeric_cols) >= 2:
                    x_var = st.sidebar.selectbox("X variable for scatter", numeric_cols, index=0, key=f"scatter_x{group_label}")
                    y_var = st.sidebar.selectbox("Y variable for scatter", numeric_cols, index=1 if len(numeric_cols) > 1 else 0, key=f"scatter_y{group_label}")
                    fig, ax = plt.subplots()
                    sns.scatterplot(data=sub_df, x=x_var, y=y_var, hue='gender_analysis', palette=BRAND_COLORS[:2])
                    plt.title(f"Scatter Plot of {y_var} vs {x_var} by Gender{group_label}")
                    st.pyplot(fig)
                    fig.savefig(buf, format='png')
                    plot_images.append((f"Scatter Plot of {y_var} vs {x_var} by Gender{group_label}", buf.getvalue()))
                    st.download_button(f"Download Scatter Plot{group_label}", buf.getvalue(), file_name=f"scatter_plot{group_label}.png", mime="image/png")
                    plot_narratives.append(f"The scatter plot visualizes the relationship between {x_var} and {y_var}, colored by gender{group_label}. Patterns or clusters may reveal gender-based trends.")
                else:
                    st.warning("Not enough numeric variables for a scatter plot.")
            if plot_type == "Pie Chart":
                pie_data = sub_df['gender_analysis'].value_counts()
                fig, ax = plt.subplots()
                ax.pie(pie_data, labels=pie_data.index, autopct='%1.1f%%', colors=BRAND_COLORS[:len(pie_data)])
                plt.title(f"Gender Distribution{group_label}")
                st.pyplot(fig)
                fig.savefig(buf, format='png')
                plot_images.append((f"Gender Distribution Pie Chart{group_label}", buf.getvalue()))
                st.download_button(f"Download Pie Chart{group_label}", buf.getvalue(), file_name=f"pie_chart{group_label}.png", mime="image/png")
                plot_narratives.append(f"The pie chart shows the proportion of males and females in your dataset{group_label}.")
            if plot_type == "Correlation Heatmap":
                numeric_cols = [col for col in columns if np.issubdtype(df[col].dropna().dtype, np.number)]
                if len(numeric_cols) > 1:
                    corr = sub_df[numeric_cols].corr()
                    fig, ax = plt.subplots()
                    sns.heatmap(corr, annot=True, cmap="Blues")
                    plt.title(f"Correlation Heatmap (Numeric Variables){group_label}")
                    st.pyplot(fig)
                    fig.savefig(buf, format='png')
                    plot_images.append((f"Correlation Heatmap{group_label}", buf.getvalue()))
                    st.download_button(f"Download Heatmap{group_label}", buf.getvalue(), file_name=f"heatmap{group_label}.png", mime="image/png")
                    plot_narratives.append(f"The correlation heatmap shows relationships between numeric variables{group_label}. Strong correlations may differ by gender.")
                else:
                    st.warning("Not enough numeric variables for a heatmap.")

    if group_col != "None":
        st.subheader(f"Breakdown by {group_col}")
        for group_val in gendered_df[group_col].dropna().unique():
            st.markdown(f"**{group_col}: {group_val}**")
            sub_df = gendered_df[gendered_df[group_col] == group_val]
            plot_for_group(sub_df, group_val)
    else:
        plot_for_group(gendered_df)

    # --- STATISTICAL TESTS & NARRATIVE ---
    st.subheader("Statistical Test & Narrative")
    if outcome_type == 'numeric':
        male_data = gendered_df[gendered_df['gender_analysis'] == 'Male'][outcome_col].dropna()
        female_data = gendered_df[gendered_df['gender_analysis'] == 'Female'][outcome_col].dropna()
        _, p1 = stats.normaltest(male_data) if len(male_data) > 7 else (None, 1)
        _, p2 = stats.normaltest(female_data) if len(female_data) > 7 else (None, 1)
        if p1 < 0.05 or p2 < 0.05:
            stat, pval = mannwhitneyu(male_data, female_data, alternative='two-sided')
            test_type = 'Mann-Whitney U'
        else:
            stat, pval = stats.ttest_ind(male_data, female_data)
            test_type = 't-test'
        st.write(f"Statistical Test: {test_type}")
        st.write(f"Statistic: {stat:.4f}")
        st.write(f"p-value: {pval:.4f}")
        st.markdown(f"""
        **Narrative:**
        - The violin and box plots above show the distribution and central tendency of {outcome_col} for each gender.
        - The average {outcome_col} for females is {female_data.mean():.2f}, for males is {male_data.mean():.2f}.
        - The {test_type} p-value is {pval:.4f}, indicating that the difference is {'not ' if pval > 0.05 else ''}statistically significant.
        """)
    else:
        crosstab = pd.crosstab(gendered_df['gender_analysis'], gendered_df[outcome_col])
        chi2, pval, dof, expected = chi2_contingency(crosstab)
        st.write(f"Chi-square statistic: {chi2:.4f}")
        st.write(f"p-value: {pval:.4f}")
        st.write(f"Degrees of freedom: {dof}")
        st.markdown(f"""
        **Narrative:**
        - The bar and pie charts above show the distribution of {outcome_col} by gender.
        - The chi-square p-value is {pval:.4f}, indicating that the difference is {'not ' if pval > 0.05 else ''}statistically significant.
        """)
    for n in plot_narratives:
        st.info(n)

    # --- DOWNLOAD SUMMARY TABLE ---
    st.download_button("Download Summary Table (CSV)", gendered_df.to_csv(index=False), file_name="summary_table.csv", mime="text/csv")

    # --- GENERATE REPORT ---
    st.subheader("Generate and Download Report")
    if st.button("Generate Report"):
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        if report_format == "PDF":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmpfile:
                doc = SimpleDocTemplate(tmpfile.name, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                # Logo
                story.append(RLImage(LOGO_PATH, width=2*inch, height=2*inch))
                # Title
                story.append(Paragraph(f"<b>{ORG_NAME} Gender Disaggregated Analysis Report</b>", styles['Title']))
                story.append(Paragraph(f"Generated: {now}", styles['Normal']))
                story.append(Spacer(1, 12))
                # Selections
                story.append(Paragraph(f"<b>Outcome:</b> {outcome_col}", styles['Normal']))
                story.append(Paragraph(f"<b>Gender column:</b> {gender_col}", styles['Normal']))
                story.append(Paragraph(f"<b>Group by:</b> {group_col}", styles['Normal']))
                story.append(Spacer(1, 12))
                # Plots
                for title, img_bytes in plot_images:
                    story.append(Paragraph(f"<b>{title}</b>", styles['Heading2']))
                    img_buf = io.BytesIO(img_bytes)
                    story.append(RLImage(img_buf, width=5*inch, height=3*inch))
                    story.append(Spacer(1, 12))
                # Narratives
                for n in plot_narratives:
                    story.append(Paragraph(n, styles['Normal']))
                # Powered by FemAnalytica
                story.append(Spacer(1, 24))
                story.append(Paragraph(f"<b>This analysis was powered by {ORG_NAME}.</b>", styles['Normal']))
                doc.build(story)
                with open(tmpfile.name, "rb") as f:
                    st.download_button("Download PDF Report", f.read(), file_name="FemAnalytica_Gender_Report.pdf", mime="application/pdf")
        else:  # Word
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmpfile:
                doc = Document()
                # Logo
                doc.add_picture(LOGO_PATH, width=Inches(2))
                doc.add_heading(f"{ORG_NAME} Gender Disaggregated Analysis Report", 0)
                doc.add_paragraph(f"Generated: {now}")
                doc.add_paragraph(f"Outcome: {outcome_col}")
                doc.add_paragraph(f"Gender column: {gender_col}")
                doc.add_paragraph(f"Group by: {group_col}")
                for title, img_bytes in plot_images:
                    doc.add_heading(title, level=2)
                    img_buf = io.BytesIO(img_bytes)
                    img_path = os.path.join(tempfile.gettempdir(), f"plot_{title.replace(' ','_')}.png")
                    with open(img_path, "wb") as f:
                        f.write(img_bytes)
                    doc.add_picture(img_path, width=Inches(5))
                for n in plot_narratives:
                    doc.add_paragraph(n)
                # Powered by FemAnalytica
                doc.add_paragraph("")
                doc.add_paragraph(f"This analysis was powered by {ORG_NAME}.")
                doc.save(tmpfile.name)
                with open(tmpfile.name, "rb") as f:
                    st.download_button("Download Word Report", f.read(), file_name="FemAnalytica_Gender_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("Please upload a CSV file to begin analysis.")

# --- FOOTER ---
st.markdown(
    f"""<hr style='border:1px solid {BRAND_COLORS[0]};'>\n
    <div style='color:{BRAND_COLORS[2]};'><b>Powered by {ORG_NAME}</b>. For support, contact info@femanalytica.org</div>
    """, unsafe_allow_html=True
) 